#!/usr/bin/env python3
"""
البوت الأكاديمي الذكي المحسن - الإصدار الشامل
جميع الميزات في ملف واحد
"""

import logging
import os
import re
import json
import uuid
import asyncio
import aiohttp
import requests
import time
import threading
from datetime import datetime, timedelta
from functools import wraps
from concurrent.futures import ThreadPoolExecutor
import io

# Telegram imports
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove, InputFile
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    ConversationHandler,
)

# Document processing imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# PDF processing imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT

# Search imports
from googlesearch import search

# --- Configuration ---
TELEGRAM_BOT_TOKEN = "7229014493:AAEOQZr_jWTkUqZmrhEfwZD69ogtL7buB8Q"
OPENROUTER_API_KEY = "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"
ADMIN_EMAIL = "java88449@gmail.com"
ADMIN_PASSWORD = "221369"
USERS_DB_FILE = "top.json"

# --- Logging Setup ---
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logging.getLogger("httpx").setLevel(logging.WARNING)
logger = logging.getLogger(__name__)

# --- Conversation States ---
(SELECTING_ACTION, WAITING_FOR_RESEARCH_INFO, GENERAL_CHAT, 
 WAITING_FOR_IMAGE_PROMPT, WAITING_FOR_TRANSLATE_TEXT,
 ADMIN_PANEL, WAITING_FOR_PASSWORD, WAITING_FOR_USERNAME, 
 WAITING_FOR_CREDITS, WAITING_FOR_TOKEN, WAITING_FOR_WEB_SEARCH) = range(11)

# =============================================================================
# DATABASE MANAGEMENT CLASS
# =============================================================================

class UserDatabase:
    """مدير قاعدة البيانات المتقدم"""
    
    def __init__(self, db_file):
        self.db_file = db_file
        self.backup_dir = "backups"
        self.load_database()
        self.ensure_backup_directory()
    
    def load_database(self):
        """تحميل قاعدة البيانات من ملف JSON"""
        try:
            with open(self.db_file, 'r', encoding='utf-8') as f:
                self.data = json.load(f)
            self.validate_database_structure()
        except FileNotFoundError:
            logger.warning(f"Database file {self.db_file} not found. Creating new database.")
            self.create_new_database()
        except json.JSONDecodeError as e:
            logger.error(f"Database corruption detected: {e}")
            self.create_new_database()
    
    def create_new_database(self):
        """إنشاء قاعدة بيانات جديدة"""
        self.data = {
            "users": {},
            "tokens": {},
            "admin_users": [ADMIN_EMAIL],
            "statistics": {
                "total_users": 0,
                "total_researches": 0,
                "total_searches": 0,
                "total_translations": 0,
                "total_tokens_generated": 0,
                "total_tokens_used": 0,
                "bot_start_date": datetime.now().isoformat(),
                "daily_stats": {},
                "monthly_stats": {}
            },
            "settings": {
                "default_credits": 0,
                "admin_credits": 999999,
                "max_research_length": 6000,
                "max_search_results": 5,
                "token_length": 18,
                "backup_frequency": 24,
                "auto_cleanup_days": 30
            },
            "logs": {
                "last_backup": None,
                "last_maintenance": None,
                "error_count": 0,
                "last_error": None
            },
            "features": {
                "research_enabled": True,
                "search_enabled": True,
                "translation_enabled": True,
                "image_search_enabled": True,
                "web_search_enabled": True
            }
        }
        self.save_database()
    
    def validate_database_structure(self):
        """التحقق من سلامة هيكل قاعدة البيانات"""
        required_keys = ["users", "tokens", "admin_users", "statistics", "settings", "logs"]
        for key in required_keys:
            if key not in self.data:
                logger.warning(f"Missing key {key} in database. Adding default value.")
                if key == "users":
                    self.data[key] = {}
                elif key == "tokens":
                    self.data[key] = {}
                elif key == "admin_users":
                    self.data[key] = [ADMIN_EMAIL]
                elif key == "statistics":
                    self.data[key] = {
                        "total_users": 0,
                        "total_researches": 0,
                        "total_searches": 0,
                        "total_translations": 0
                    }
                elif key == "settings":
                    self.data[key] = {
                        "default_credits": 0,
                        "admin_credits": 999999
                    }
                elif key == "logs":
                    self.data[key] = {
                        "last_backup": None,
                        "last_maintenance": None,
                        "error_count": 0
                    }
    
    def ensure_backup_directory(self):
        """التأكد من وجود مجلد النسخ الاحتياطية"""
        if not os.path.exists(self.backup_dir):
            os.makedirs(self.backup_dir)
    
    def save_database(self):
        """حفظ قاعدة البيانات مع نسخ احتياطي"""
        try:
            # إنشاء نسخة احتياطية قبل الحفظ
            if os.path.exists(self.db_file):
                backup_name = f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                backup_path = os.path.join(self.backup_dir, backup_name)
                try:
                    with open(self.db_file, 'r', encoding='utf-8') as src:
                        with open(backup_path, 'w', encoding='utf-8') as dst:
                            dst.write(src.read())
                except:
                    pass  # تجاهل أخطاء النسخ الاحتياطي
            
            # حفظ البيانات الجديدة
            with open(self.db_file, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, ensure_ascii=False, indent=2)
            
            self.data["logs"]["last_backup"] = datetime.now().isoformat()
            
        except Exception as e:
            logger.error(f"Error saving database: {e}")
            self.data["logs"]["error_count"] += 1
            self.data["logs"]["last_error"] = str(e)
    
    def add_user(self, user_id, username, first_name=None, last_name=None, credits=0):
        """إضافة مستخدم جديد"""
        user_id = str(user_id)
        if user_id not in self.data["users"]:
            self.data["users"][user_id] = {
                "username": username,
                "first_name": first_name,
                "last_name": last_name,
                "credits": credits,
                "join_date": datetime.now().isoformat(),
                "last_activity": datetime.now().isoformat(),
                "usage_count": 0,
                "research_count": 0,
                "search_count": 0,
                "translation_count": 0,
                "total_spent_credits": 0,
                "status": "active"
            }
            self.data["statistics"]["total_users"] += 1
            self.update_daily_stats("new_users", 1)
            self.save_database()
    
    def get_user(self, user_id):
        """الحصول على بيانات المستخدم"""
        return self.data["users"].get(str(user_id))
    
    def update_user_credits(self, user_id, credits):
        """تحديث رصيد المستخدم"""
        user_id = str(user_id)
        if user_id in self.data["users"]:
            self.data["users"][user_id]["credits"] = credits
            self.save_database()
            return True
        return False
    
    def use_credit(self, user_id):
        """استخدام رصيد واحد"""
        user_id = str(user_id)
        if user_id in self.data["users"] and self.data["users"][user_id]["credits"] > 0:
            self.data["users"][user_id]["credits"] -= 1
            self.data["users"][user_id]["usage_count"] += 1
            self.data["users"][user_id]["last_activity"] = datetime.now().isoformat()
            self.save_database()
            return True
        return False
    
    def update_user_activity(self, user_id, activity_type):
        """تحديث نشاط المستخدم"""
        user_id = str(user_id)
        if user_id in self.data["users"]:
            self.data["users"][user_id]["last_activity"] = datetime.now().isoformat()
            self.data["users"][user_id]["usage_count"] += 1
            
            if activity_type == "research":
                self.data["users"][user_id]["research_count"] += 1
                self.data["statistics"]["total_researches"] += 1
                self.update_daily_stats("researches", 1)
            elif activity_type == "search":
                self.data["users"][user_id]["search_count"] += 1
                self.data["statistics"]["total_searches"] += 1
                self.update_daily_stats("searches", 1)
            elif activity_type == "translation":
                self.data["users"][user_id]["translation_count"] += 1
                self.data["statistics"]["total_translations"] += 1
                self.update_daily_stats("translations", 1)
            
            self.save_database()
    
    def update_daily_stats(self, stat_type, increment=1):
        """تحديث الإحصائيات اليومية"""
        today = datetime.now().strftime('%Y-%m-%d')
        if "daily_stats" not in self.data["statistics"]:
            self.data["statistics"]["daily_stats"] = {}
        
        if today not in self.data["statistics"]["daily_stats"]:
            self.data["statistics"]["daily_stats"][today] = {
                "new_users": 0,
                "researches": 0,
                "searches": 0,
                "translations": 0,
                "tokens_used": 0
            }
        
        if stat_type in self.data["statistics"]["daily_stats"][today]:
            self.data["statistics"]["daily_stats"][today][stat_type] += increment
    
    def generate_token(self, credits, expiry_days=30, description=""):
        """توليد توكن متقدم"""
        token = str(uuid.uuid4()).replace('-', '')[:18].upper()
        expiry_date = datetime.now() + timedelta(days=expiry_days)
        
        self.data["tokens"][token] = {
            "credits": credits,
            "created_date": datetime.now().isoformat(),
            "expiry_date": expiry_date.isoformat(),
            "description": description,
            "used": False,
            "used_by": None,
            "used_date": None
        }
        
        self.data["statistics"]["total_tokens_generated"] += 1
        self.save_database()
        return token
    
    def use_token(self, user_id, token):
        """استخدام التوكن"""
        if token in self.data["tokens"]:
            token_data = self.data["tokens"][token]
            
            # التحقق من الاستخدام السابق
            if token_data["used"]:
                return {"success": False, "message": "التوكن مستخدم من قبل"}
            
            # التحقق من تاريخ الانتهاء
            if "expiry_date" in token_data:
                expiry = datetime.fromisoformat(token_data["expiry_date"])
                if datetime.now() > expiry:
                    return {"success": False, "message": "التوكن منتهي الصلاحية"}
            
            # استخدام التوكن
            credits = token_data["credits"]
            token_data["used"] = True
            token_data["used_by"] = str(user_id)
            token_data["used_date"] = datetime.now().isoformat()
            
            # إضافة الرصيد للمستخدم
            user_id = str(user_id)
            if user_id in self.data["users"]:
                self.data["users"][user_id]["credits"] += credits
            
            self.data["statistics"]["total_tokens_used"] += 1
            self.update_daily_stats("tokens_used", 1)
            self.save_database()
            
            return {"success": True, "credits": credits, "message": "تم تفعيل التوكن بنجاح"}
        
        return {"success": False, "message": "التوكن غير صحيح"}
    
    def get_system_statistics(self):
        """الحصول على إحصائيات النظام"""
        stats = self.data["statistics"]
        
        # حساب المستخدمين النشطين (آخر 7 أيام)
        week_ago = datetime.now() - timedelta(days=7)
        active_users = 0
        for user in self.data["users"].values():
            if user.get("last_activity"):
                try:
                    last_activity = datetime.fromisoformat(user["last_activity"])
                    if last_activity > week_ago:
                        active_users += 1
                except:
                    pass
        
        return {
            "total_users": stats.get("total_users", 0),
            "active_users_week": active_users,
            "total_researches": stats.get("total_researches", 0),
            "total_searches": stats.get("total_searches", 0),
            "total_translations": stats.get("total_translations", 0),
            "total_tokens_generated": stats.get("total_tokens_generated", 0),
            "total_tokens_used": stats.get("total_tokens_used", 0),
            "bot_start_date": stats.get("bot_start_date"),
            "database_size": self.get_database_size()
        }
    
    def get_database_size(self):
        """الحصول على حجم قاعدة البيانات"""
        try:
            size = os.path.getsize(self.db_file)
            if size < 1024:
                return f"{size} B"
            elif size < 1024 * 1024:
                return f"{size / 1024:.1f} KB"
            else:
                return f"{size / (1024 * 1024):.1f} MB"
        except:
            return "Unknown"
    
    def is_admin(self, email):
        """التحقق من صلاحيات الإدارة"""
        return email in self.data["admin_users"]

# Initialize database
db = UserDatabase(USERS_DB_FILE)

# =============================================================================
# PERFORMANCE OPTIMIZATION CLASS
# =============================================================================

class PerformanceOptimizer:
    """فئة تحسين الأداء"""
    
    def __init__(self):
        self.cache = {}
        self.cache_ttl = {}
        self.session_pool = None
        self.response_times = []
        
    async def initialize_session_pool(self):
        """تهيئة مجموعة الجلسات للطلبات المتزامنة"""
        connector = aiohttp.TCPConnector(
            limit=100,
            limit_per_host=30,
            ttl_dns_cache=300,
            use_dns_cache=True,
        )
        timeout = aiohttp.ClientTimeout(total=30, connect=10)
        self.session_pool = aiohttp.ClientSession(
            connector=connector,
            timeout=timeout
        )
    
    async def close_session_pool(self):
        """إغلاق مجموعة الجلسات"""
        if self.session_pool:
            await self.session_pool.close()
    
    def cache_response(self, ttl=300):
        """ديكوريتر لتخزين الاستجابات مؤقتاً"""
        def decorator(func):
            @wraps(func)
            async def wrapper(*args, **kwargs):
                # إنشاء مفتاح التخزين المؤقت
                cache_key = f"{func.__name__}_{hash(str(args) + str(kwargs))}"
                
                # التحقق من وجود البيانات في التخزين المؤقت
                if cache_key in self.cache:
                    if time.time() - self.cache_ttl[cache_key] < ttl:
                        logger.debug(f"Cache hit for {func.__name__}")
                        return self.cache[cache_key]
                    else:
                        # إزالة البيانات المنتهية الصلاحية
                        del self.cache[cache_key]
                        del self.cache_ttl[cache_key]
                
                # تنفيذ الوظيفة وتخزين النتيجة
                result = await func(*args, **kwargs)
                self.cache[cache_key] = result
                self.cache_ttl[cache_key] = time.time()
                logger.debug(f"Cache miss for {func.__name__}, result cached")
                
                return result
            return wrapper
        return decorator
    
    async def optimized_api_call(self, url, headers, payload, max_retries=3):
        """استدعاء API محسن مع إعادة المحاولة"""
        if not self.session_pool:
            await self.initialize_session_pool()
        
        for attempt in range(max_retries):
            try:
                async with self.session_pool.post(url, headers=headers, json=payload) as response:
                    if response.status == 200:
                        return await response.json()
                    elif response.status == 429:  # Rate limit
                        wait_time = 2 ** attempt
                        logger.warning(f"Rate limited, waiting {wait_time}s before retry")
                        await asyncio.sleep(wait_time)
                    else:
                        logger.error(f"API error: {response.status}")
                        if attempt == max_retries - 1:
                            return None
            except asyncio.TimeoutError:
                logger.warning(f"Timeout on attempt {attempt + 1}")
                if attempt == max_retries - 1:
                    return None
            except Exception as e:
                logger.error(f"API call error on attempt {attempt + 1}: {e}")
                if attempt == max_retries - 1:
                    return None
        
        return None

# Initialize performance optimizer
performance_optimizer = PerformanceOptimizer()

# =============================================================================
# RATE LIMITER CLASS
# =============================================================================

class RateLimiter:
    """محدد معدل الطلبات"""
    
    def __init__(self, max_requests=10, time_window=60):
        self.max_requests = max_requests
        self.time_window = time_window
        self.requests = {}
        self.lock = threading.Lock()
    
    def is_allowed(self, user_id):
        """التحقق من السماح بالطلب"""
        current_time = time.time()
        user_id = str(user_id)
        
        with self.lock:
            if user_id not in self.requests:
                self.requests[user_id] = []
            
            # إزالة الطلبات القديمة
            self.requests[user_id] = [
                req_time for req_time in self.requests[user_id]
                if current_time - req_time < self.time_window
            ]
            
            # التحقق من الحد الأقصى
            if len(self.requests[user_id]) >= self.max_requests:
                return False
            
            # إضافة الطلب الحالي
            self.requests[user_id].append(current_time)
            return True

# Initialize rate limiter
rate_limiter = RateLimiter(max_requests=20, time_window=60)

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def is_rtl(text: str) -> bool:
    """التحقق من النص العربي"""
    rtl_chars = re.findall(r'[\u0590-\u05FF\u0600-\u06FF]', text)
    return len(rtl_chars) > len(text) / 4

def set_rtl_paragraph(paragraph):
    """تعيين خصائص الفقرة للنص العربي"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement('w:bidi'))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = 'Arial'
    run.font.rtl = True

async def call_openrouter_api(prompt: str, model="deepseek/deepseek-chat", max_tokens=4000) -> str:
    """استدعاء API محسن مع معالجة أفضل للأخطاء"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": 0.7
    }
    
    url = "https://openrouter.ai/api/v1/chat/completions"
    result = await performance_optimizer.optimized_api_call(url, headers, payload)
    
    if result and "choices" in result and len(result["choices"]) > 0:
        return result["choices"][0]["message"]["content"]
    
    return None

async def search_web(query: str, num_results: int = 5) -> list:
    """البحث في الإنترنت مع معالجة محسنة"""
    try:
        results = []
        for url in search(query, num_results=num_results, lang="ar"):
            results.append(url)
        return results
    except Exception as e:
        logger.error(f"Web search error: {e}")
        return []

async def search_images(query: str, num_results: int = 3) -> list:
    """البحث عن الصور"""
    try:
        results = []
        search_generator = search(query, tbm='isch', num_results=num_results, lang="ar")
        for url in search_generator:
            results.append(url)
        return results
    except Exception as e:
        logger.error(f"Image search error: {e}")
        return []

# =============================================================================
# DOCUMENT CREATION FUNCTIONS
# =============================================================================

async def create_enhanced_word_document(title, university, department, stage, subject, 
                                      student_name, section, doctor_name, content):
    """إنشاء مستند Word محسن بتنسيق جميل"""
    document = Document()
    
    # تعيين اتجاه المستند للعربية
    doc_element = document.element.body.sectPr
    doc_element.append(OxmlElement('w:bidi'))
    
    # إنشاء أنماط مخصصة
    styles = document.styles
    
    # نمط العنوان الرئيسي
    try:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # أزرق داكن
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    except:
        pass  # تجاهل أخطاء الأنماط
    
    # صفحة الغلاف
    # شعار الجامعة (مساحة فارغة)
    p = document.add_paragraph()
    p.add_run('\n\n')
    
    # اسم الجامعة
    uni_para = document.add_paragraph()
    uni_run = uni_para.add_run(university)
    uni_run.font.name = 'Arial'
    uni_run.font.size = Pt(22)
    uni_run.font.bold = True
    uni_run.font.color.rgb = RGBColor(0, 51, 102)
    uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # القسم
    dept_para = document.add_paragraph()
    dept_run = dept_para.add_run(f"كلية {department}")
    dept_run.font.name = 'Arial'
    dept_run.font.size = Pt(16)
    dept_run.font.bold = True
    dept_run.font.color.rgb = RGBColor(0, 102, 204)
    dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # مساحة
    document.add_paragraph('\n\n')
    
    # عنوان البحث
    title_para = document.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('\n\n')
    
    # جدول معلومات الطالب
    table = document.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    info_data = [
        ("إعداد الطالب:", student_name),
        ("المرحلة:", f"{stage} / الشعبة: {section}"),
        ("المادة:", subject),
        ("بإشراف الدكتور:", doctor_name),
        ("العام الدراسي:", "2024-2025")
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        
        # تنسيق الخلايا
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
    
    # كسر صفحة
    document.add_page_break()
    
    # محتوى البحث
    content_lines = content.split('\n')
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        
        p = document.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        
        # تحديد نوع النص (عنوان أم فقرة)
        if any(keyword in line for keyword in ['مقدمة', 'خاتمة', 'المراجع', 'الفصل', 'المبحث']):
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # تعيين خصائص RTL
        if is_rtl(line):
            set_rtl_paragraph(p)
    
    # حفظ الملف
    filename = f"research_{title.replace(' ', '_')}.docx"
    document.save(filename)
    return filename

async def create_enhanced_pdf_document(title, university, department, stage, subject,
                                     student_name, section, doctor_name, content):
    """إنشاء مستند PDF محسن بتنسيق جميل"""
    filename = f"research_{title.replace(' ', '_')}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=A4)
    
    # تحديد الأنماط
    styles = getSampleStyleSheet()
    
    # نمط العنوان الرئيسي
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=20,
        textColor=colors.HexColor('#003366'),
        alignment=TA_CENTER,
        spaceAfter=20,
        fontName='Helvetica-Bold'
    )
    
    # نمط العناوين الفرعية
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#0066CC'),
        alignment=TA_RIGHT,
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    
    # نمط النص العادي
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_RIGHT,
        spaceBefore=6,
        spaceAfter=6,
        fontName='Helvetica'
    )
    
    # بناء المحتوى
    story = []
    
    # صفحة الغلاف
    story.append(Spacer(1, 50))
    story.append(Paragraph(university, title_style))
    story.append(Paragraph(f"كلية {department}", heading_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 40))
    
    # جدول معلومات الطالب
    info_data = [
        ["إعداد الطالب:", student_name],
        ["المرحلة:", f"{stage} / الشعبة: {section}"],
        ["المادة:", subject],
        ["بإشراف الدكتور:", doctor_name],
        ["العام الدراسي:", "2024-2025"]
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 3*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F0F8FF')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E6F3FF')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#0066CC'))
    ]))
    
    story.append(info_table)
    story.append(PageBreak())
    
    # محتوى البحث
    content_lines = content.split('\n')
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        
        # تحديد نوع النص
        if any(keyword in line for keyword in ['مقدمة', 'خاتمة', 'المراجع', 'الفصل', 'المبحث']):
            story.append(Paragraph(line, heading_style))
        else:
            story.append(Paragraph(line, normal_style))
    
    # بناء المستند
    doc.build(story)
    return filename

# =============================================================================
# USER AUTHENTICATION
# =============================================================================

async def check_user_access(update: Update) -> bool:
    """التحقق من صلاحية المستخدم"""
    user_id = update.effective_user.id
    username = update.effective_user.username or update.effective_user.first_name
    
    # إضافة المستخدم إذا لم يكن موجوداً
    if not db.get_user(user_id):
        db.add_user(user_id, username, update.effective_user.first_name, update.effective_user.last_name, 0)
    
    # التحقق من الإدارة
    if update.message.text == ADMIN_EMAIL:
        return True
    
    # التحقق من الرصيد
    user = db.get_user(user_id)
    return user and (user["credits"] > 0 or str(update.effective_user.username) == ADMIN_EMAIL.split('@')[0])

# =============================================================================
# MENU KEYBOARDS
# =============================================================================

main_menu_keyboard = [
    ["انشاء بحث جامعي 🎓"],
    ["استخدم البوت (سؤال وجواب) 🤖"],
    ["بحث عن صورة 🖼️", "بحث في الإنترنت 🌐"],
    ["ترجمه 🌍", "استخدام توكن 🎫"],
]

admin_menu_keyboard = [
    ["لوحة التحكم 🔧"],
    ["العودة للقائمة الرئيسية 🏠"]
]

# =============================================================================
# MAIN COMMAND HANDLERS
# =============================================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء المحادثة وعرض القائمة الرئيسية"""
    user = update.effective_user
    user_id = user.id
    username = user.username or user.first_name
    
    # إضافة المستخدم إذا لم يكن موجوداً
    if not db.get_user(user_id):
        db.add_user(user_id, username, user.first_name, user.last_name, 0)
    
    # التحقق من الإدارة
    if update.message.text == ADMIN_EMAIL:
        keyboard = main_menu_keyboard + admin_menu_keyboard
        markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True, resize_keyboard=True)
        await update.message.reply_html(
            f"أهلاً بك {user.mention_html()}! 👋\n\n🔧 **وضع الإدارة مُفعل**\n\nأنا مساعدك الأكاديمي الشامل. اختر إحدى الخدمات:",
            reply_markup=markup,
        )
    else:
        user_data = db.get_user(user_id)
        credits = user_data["credits"] if user_data else 0
        
        markup = ReplyKeyboardMarkup(main_menu_keyboard, one_time_keyboard=True, resize_keyboard=True)
        await update.message.reply_html(
            f"أهلاً بك {user.mention_html()}! 👋\n\n💰 **رصيدك الحالي: {credits} استخدام**\n\nأنا مساعدك الأكاديمي الشامل. اختر إحدى الخدمات:",
            reply_markup=markup,
        )
    
    return SELECTING_ACTION

# =============================================================================
# RESEARCH FEATURE
# =============================================================================

async def start_research(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """طلب معلومات البحث الأكاديمي"""
    if not await check_user_access(update):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    message = """🎓 **إنشاء بحث جامعي**

يرجى إرسال جميع المعلومات التالية:

1. عنوان السيمنار :

2. اسم الجامعة :

3. القسم :

4. المرحلة :

5. المادة :

6. الاسم: 

7. الشعبة: 

8. اسم الدكتور: 

10. تفاصيل إضافية عن السيمنار :"""
    
    await update.message.reply_text(message)
    return WAITING_FOR_RESEARCH_INFO

async def process_research_info(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة معلومات البحث وإنشاء الملفات"""
    user_id = update.effective_user.id
    
    # استخدام رصيد
    if not db.use_credit(user_id):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي.")
        return SELECTING_ACTION
    
    await update.message.reply_text("جاري إنشاء البحث تبعك يرجى الانتظار بصبر")
    await update.message.reply_text("🚀")
    
    user_input = update.message.text
    lines = user_input.strip().split("\n")
    user_data = {}
    
    for line in lines:
        if ":" in line:
            key, value = line.split(":", 1)
            user_data[key.strip()] = value.strip()
    
    # استخراج البيانات
    seminar_title = user_data.get('عنوان السيمنار', 'بحث أكاديمي')
    university = user_data.get('اسم الجامعة', '')
    department = user_data.get('القسم', '')
    stage = user_data.get('المرحلة', '')
    subject = user_data.get('المادة', '')
    student_name = user_data.get('الاسم', '')
    section = user_data.get('الشعبة', '')
    doctor_name = user_data.get('اسم الدكتور', '')
    details = user_data.get('تفاصيل إضافية عن السيمنار', '')
    
    # إنشاء محتوى البحث
    prompt = f"""
    أنت باحث أكاديمي متخصص. اكتب بحثاً شاملاً ومفصلاً باللغة العربية حول الموضوع: '{seminar_title}'.
    التفاصيل الإضافية: '{details}'.
    
    يجب أن يتضمن البحث:
    1. مقدمة شاملة ومفصلة (500 كلمة على الأقل)
    2. محتوى رئيسي مقسم إلى 4-5 أقسام رئيسية مع عناوين فرعية واضحة
    3. كل قسم يجب أن يحتوي على 800-1000 كلمة
    4. أمثلة وتطبيقات عملية
    5. خاتمة شاملة تلخص النتائج الرئيسية (400 كلمة)
    6. قائمة بالمراجع والمصادر (10 مراجع على الأقل)
    
    اكتب المحتوى بأسلوب أكاديمي ومفصل جداً. استخدم لغة علمية دقيقة.
    """
    
    research_content = await call_openrouter_api(prompt, max_tokens=6000)
    if not research_content:
        await update.message.reply_text("عذراً، حدث خطأ أثناء توليد المحتوى.")
        return SELECTING_ACTION
    
    try:
        # إنشاء ملف Word
        word_file = await create_enhanced_word_document(
            seminar_title, university, department, stage, subject,
            student_name, section, doctor_name, research_content
        )
        
        # إنشاء ملف PDF
        pdf_file = await create_enhanced_pdf_document(
            seminar_title, university, department, stage, subject,
            student_name, section, doctor_name, research_content
        )
        
        # إرسال الملفات
        await update.message.reply_document(
            document=InputFile(word_file),
            filename=f"{seminar_title.replace(' ', '_')}.docx",
            caption="📄 ملف Word - بحثك جاهز بتنسيق احترافي!"
        )
        
        await update.message.reply_document(
            document=InputFile(pdf_file),
            filename=f"{seminar_title.replace(' ', '_')}.pdf",
            caption="📋 ملف PDF - نسخة للطباعة والعرض!"
        )
        
        # حذف الملفات المؤقتة
        try:
            os.unlink(word_file)
            os.unlink(pdf_file)
        except:
            pass
        
        # تحديث الإحصائيات
        db.update_user_activity(user_id, "research")
        
    except Exception as e:
        logger.error(f"Error creating/sending files: {e}")
        await update.message.reply_text("حدث خطأ أثناء إنشاء أو إرسال الملفات.")
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# GENERAL CHAT FEATURE
# =============================================================================

async def start_general_chat(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء وضع السؤال والجواب"""
    if not await check_user_access(update):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    await update.message.reply_text(
        "🤖 **وضع السؤال والجواب**\n\nاسألني أي شيء وسأجيبك باستخدام الذكاء الاصطناعي.\n\nللخروج والعودة للقائمة، أرسل /cancel.",
        reply_markup=ReplyKeyboardRemove()
    )
    return GENERAL_CHAT

async def handle_general_chat(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة الرسائل في وضع المحادثة العامة"""
    user_id = update.effective_user.id
    
    # استخدام رصيد
    if not db.use_credit(user_id):
        await update.message.reply_text("❌ عذراً، انتهى رصيدك. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    question = update.message.text
    await update.message.reply_text("🤔 أفكر في إجابة...")
    
    answer = await call_openrouter_api(question)
    if answer:
        await update.message.reply_text(answer)
        db.update_user_activity(user_id, "search")
    else:
        await update.message.reply_text("عذراً، لم أتمكن من معالجة طلبك.")
    
    return GENERAL_CHAT

# =============================================================================
# IMAGE SEARCH FEATURE
# =============================================================================

async def start_image_search(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء البحث عن الصور"""
    if not await check_user_access(update):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    await update.message.reply_text("🖼️ **بحث عن صورة**\n\nأرسل لي وصفاً للصورة التي تريد البحث عنها.")
    return WAITING_FOR_IMAGE_PROMPT

async def handle_image_search(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة البحث عن الصور"""
    user_id = update.effective_user.id
    
    # استخدام رصيد
    if not db.use_credit(user_id):
        await update.message.reply_text("❌ عذراً، انتهى رصيدك.")
        return SELECTING_ACTION
    
    query = update.message.text
    await update.message.reply_text(f"🔍 أبحث عن صور لـ '{query}'...")
    
    try:
        images = await search_images(query, 3)
        if images:
            await update.message.reply_text(f"✅ وجدت {len(images)} صورة:")
            for i, image_url in enumerate(images, 1):
                try:
                    await update.message.reply_photo(
                        photo=image_url, 
                        caption=f"صورة {i} - {query}"
                    )
                except:
                    continue
        else:
            await update.message.reply_text("❌ لم أتمكن من العثور على صور مناسبة.")
        
        # تحديث الإحصائيات
        db.update_user_activity(user_id, "search")
        
    except Exception as e:
        logger.error(f"Image search error: {e}")
        await update.message.reply_text("حدث خطأ أثناء البحث عن الصور.")
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# WEB SEARCH FEATURE
# =============================================================================

async def start_web_search(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء البحث في الإنترنت"""
    if not await check_user_access(update):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    await update.message.reply_text("🌐 **البحث في الإنترنت**\n\nأرسل لي ما تريد البحث عنه.")
    return WAITING_FOR_WEB_SEARCH

async def handle_web_search(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة البحث في الإنترنت"""
    user_id = update.effective_user.id
    
    # استخدام رصيد
    if not db.use_credit(user_id):
        await update.message.reply_text("❌ عذراً، انتهى رصيدك.")
        return SELECTING_ACTION
    
    query = update.message.text
    await update.message.reply_text(f"🔍 أبحث في الإنترنت عن '{query}'...")
    
    try:
        results = await search_web(query, 5)
        if results:
            response = f"✅ **نتائج البحث عن '{query}':**\n\n"
            for i, url in enumerate(results, 1):
                response += f"{i}. {url}\n\n"
            await update.message.reply_text(response)
        else:
            await update.message.reply_text("❌ لم أتمكن من العثور على نتائج.")
        
        # تحديث الإحصائيات
        db.update_user_activity(user_id, "search")
        
    except Exception as e:
        logger.error(f"Web search error: {e}")
        await update.message.reply_text("حدث خطأ أثناء البحث.")
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# TRANSLATION FEATURE
# =============================================================================

async def start_translation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء خدمة الترجمة"""
    if not await check_user_access(update):
        await update.message.reply_text("❌ عذراً، ليس لديك رصيد كافي. استخدم توكن أو تواصل مع الإدارة.")
        return SELECTING_ACTION
    
    await update.message.reply_text("🌍 **خدمة الترجمة**\n\nأرسل لي النص الذي تريد ترجمته.")
    return WAITING_FOR_TRANSLATE_TEXT

async def handle_translation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة الترجمة"""
    user_id = update.effective_user.id
    
    # استخدام رصيد
    if not db.use_credit(user_id):
        await update.message.reply_text("❌ عذراً، انتهى رصيدك.")
        return SELECTING_ACTION
    
    text_to_translate = update.message.text
    await update.message.reply_text("🔄 جارِ الترجمة...")
    
    prompt = f"ترجم النص التالي إلى اللغة المناسبة (إذا كان عربي ترجم للإنجليزية، وإذا كان إنجليزي ترجم للعربية). قدم الترجمة فقط:\n\n'{text_to_translate}'"
    
    translated_text = await call_openrouter_api(prompt, model="google/gemini-flash-1.5", max_tokens=1000)
    
    if translated_text:
        await update.message.reply_text(f"**الترجمة:**\n\n{translated_text}")
        db.update_user_activity(user_id, "translation")
    else:
        await update.message.reply_text("عذراً، حدث خطأ أثناء الترجمة.")
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# TOKEN SYSTEM
# =============================================================================

async def start_token_usage(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء استخدام التوكن"""
    await update.message.reply_text("🎫 **استخدام توكن**\n\nأرسل لي التوكن الخاص بك:")
    return WAITING_FOR_TOKEN

async def handle_token_usage(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة استخدام التوكن"""
    user_id = update.effective_user.id
    username = update.effective_user.username or update.effective_user.first_name
    token = update.message.text.strip()
    
    # إضافة المستخدم إذا لم يكن موجوداً
    if not db.get_user(user_id):
        db.add_user(user_id, username, update.effective_user.first_name, update.effective_user.last_name, 0)
    
    result = db.use_token(user_id, token)
    if result["success"]:
        await update.message.reply_text(f"✅ {result['message']}\n💰 تم إضافة {result['credits']} استخدام لحسابك.")
    else:
        await update.message.reply_text(f"❌ {result['message']}")
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# ADMIN PANEL
# =============================================================================

async def start_admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """بدء لوحة التحكم"""
    await update.message.reply_text("🔐 **لوحة التحكم**\n\nأدخل كلمة المرور:")
    return WAITING_FOR_PASSWORD

async def handle_admin_password(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة كلمة مرور الإدارة"""
    password = update.message.text.strip()
    if password == ADMIN_PASSWORD:
        await update.message.reply_text("✅ تم تسجيل الدخول بنجاح!\n\n👤 أدخل اسم المستخدم (بالإنجليزية بدون فراغات):")
        return WAITING_FOR_USERNAME
    else:
        await update.message.reply_text("❌ كلمة مرور خاطئة!")
        return SELECTING_ACTION

async def handle_admin_username(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة اسم المستخدم للإدارة"""
    username = update.message.text.strip()
    context.user_data['target_username'] = username
    await update.message.reply_text(f"💰 أدخل كمية الرصيد لإضافتها للمستخدم {username}:")
    return WAITING_FOR_CREDITS

async def handle_admin_credits(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة إضافة الرصيد"""
    try:
        credits = int(update.message.text.strip())
        username = context.user_data.get('target_username')
        
        # توليد توكن
        token = db.generate_token(credits)
        
        await update.message.reply_text(
            f"✅ تم إنشاء توكن بنجاح!\n\n"
            f"👤 المستخدم: {username}\n"
            f"💰 الرصيد: {credits}\n"
            f"🎫 التوكن: `{token}`\n\n"
            f"⚠️ هذا التوكن صالح لمرة واحدة فقط!",
            parse_mode='Markdown'
        )
        
        # إحصائيات
        stats = db.get_system_statistics()
        await update.message.reply_text(
            f"📊 **إحصائيات النظام:**\n\n"
            f"👥 إجمالي المستخدمين: {stats['total_users']}\n"
            f"📚 إجمالي الأبحاث: {stats['total_researches']}\n"
            f"🔍 إجمالي عمليات البحث: {stats['total_searches']}\n"
            f"🌍 إجمالي الترجمات: {stats['total_translations']}\n"
            f"🎫 التوكنات المُنشأة: {stats['total_tokens_generated']}\n"
            f"✅ التوكنات المُستخدمة: {stats['total_tokens_used']}"
        )
        
    except ValueError:
        await update.message.reply_text("❌ يرجى إدخال رقم صحيح!")
        return WAITING_FOR_CREDITS
    
    await update.message.reply_text("هل ترغب في خدمة أخرى؟", 
                                   reply_markup=ReplyKeyboardMarkup(main_menu_keyboard + admin_menu_keyboard, 
                                                                  one_time_keyboard=True, 
                                                                  resize_keyboard=True))
    return SELECTING_ACTION

# =============================================================================
# MESSAGE HANDLERS
# =============================================================================

async def handle_menu_selection(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """معالجة اختيارات القائمة"""
    text = update.message.text
    
    if text == "انشاء بحث جامعي 🎓":
        return await start_research(update, context)
    elif text == "استخدم البوت (سؤال وجواب) 🤖":
        return await start_general_chat(update, context)
    elif text == "بحث عن صورة 🖼️":
        return await start_image_search(update, context)
    elif text == "بحث في الإنترنت 🌐":
        return await start_web_search(update, context)
    elif text == "ترجمه 🌍":
        return await start_translation(update, context)
    elif text == "استخدام توكن 🎫":
        return await start_token_usage(update, context)
    elif text == "لوحة التحكم 🔧":
        return await start_admin_panel(update, context)
    elif text == "العودة للقائمة الرئيسية 🏠":
        return await start(update, context)
    else:
        await update.message.reply_text("يرجى اختيار خيار من القائمة.")
        return SELECTING_ACTION

# =============================================================================
# CONVERSATION FALLBACK
# =============================================================================

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """إلغاء والعودة للقائمة الرئيسية"""
    await update.message.reply_text(
        "تم الإلغاء. عدنا إلى القائمة الرئيسية.", 
        reply_markup=ReplyKeyboardMarkup(main_menu_keyboard, one_time_keyboard=True, resize_keyboard=True)
    )
    return SELECTING_ACTION

# =============================================================================
# MAIN BOT EXECUTION
# =============================================================================

async def initialize_bot():
    """تهيئة البوت"""
    try:
        await performance_optimizer.initialize_session_pool()
        logger.info("✅ تم تهيئة البوت بنجاح")
    except Exception as e:
        logger.error(f"❌ خطأ في تهيئة البوت: {e}")

async def cleanup_bot():
    """تنظيف موارد البوت"""
    try:
        await performance_optimizer.close_session_pool()
        db.save_database()
        logger.info("✅ تم تنظيف موارد البوت")
    except Exception as e:
        logger.error(f"❌ خطأ في تنظيف البوت: {e}")

def main() -> None:
    """تشغيل البوت"""
    print("🚀 بدء تشغيل البوت الأكاديمي الذكي المحسن...")
    
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            SELECTING_ACTION: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_menu_selection),
            ],
            WAITING_FOR_RESEARCH_INFO: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, process_research_info)
            ],
            GENERAL_CHAT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_general_chat)
            ],
            WAITING_FOR_IMAGE_PROMPT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_image_search)
            ],
            WAITING_FOR_WEB_SEARCH: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_web_search)
            ],
            WAITING_FOR_TRANSLATE_TEXT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_translation)
            ],
            WAITING_FOR_TOKEN: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_token_usage)
            ],
            WAITING_FOR_PASSWORD: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_admin_password)
            ],
            WAITING_FOR_USERNAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_admin_username)
            ],
            WAITING_FOR_CREDITS: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_admin_credits)
            ],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    application.add_handler(conv_handler)
    application.add_handler(CommandHandler("start", start))

    # تشغيل البوت
    try:
        print("🤖 البوت يعمل الآن...")
        application.run_polling(allowed_updates=Update.ALL_TYPES)
    except KeyboardInterrupt:
        print("\n👋 تم إيقاف البوت بواسطة المستخدم")
    except Exception as e:
        print(f"❌ خطأ في تشغيل البوت: {e}")
    finally:
        print("🔄 تنظيف الموارد...")

if __name__ == "__main__":
    main()

